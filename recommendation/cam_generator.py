import json
import os
import re
import uuid
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

try:
    from docx2pdf import convert as docx2pdf_convert
    _DOCX2PDF_AVAILABLE = True
except ImportError:
    _DOCX2PDF_AVAILABLE = False

from common.gemini_client import get_gemini_client
from recommendation.prompts import get_cam_narrative_prompt
from recommendation.schemas import (
    CreditDecision,
    FiveCsScore,
    RecommendationRequest,
    ScoreExplanation,
)

CAM_STORAGE_DIR = "storage/cam"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _table_row(table, label: str, value: str):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)


def _safe_company_name(name: str) -> str:
    return re.sub(r"[^\w\s-]", "", name).strip().replace(" ", "_") or "Company"


# ── Main generator ────────────────────────────────────────────────────────────

async def generate_cam_document(
    request: RecommendationRequest,
    score: FiveCsScore,
    score_explanation: ScoreExplanation,
    decision: CreditDecision,
) -> str:
    """Build the CAM Word document and return the saved filename."""

    # 1. Generate narrative sections via Gemini --------------------------------
    gemini = get_gemini_client()

    request_data = {
        "company_name": request.company_name,
        "sector": request.sector,
        "loan_requested": request.loan_requested,
        "collateral_value": request.collateral_value,
        "collateral_description": request.collateral_description,
        "promoters": request.promoters,
        "research_summary": request.research_summary,
        "risk_signals": [s.model_dump() for s in request.risk_signals[:6]],
        "sector_trends": request.sector_trends[:5],
        "financial_data": request.financial_data.model_dump() if request.financial_data else {},
    }

    prompt = get_cam_narrative_prompt(request_data, score.model_dump(), decision.model_dump())

    gemini_response = await gemini.aio.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config={"temperature": 0.3, "response_mime_type": "application/json"},
    )

    raw = gemini_response.text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    narrative: dict = json.loads(raw)

    # Ensure every narrative value is a plain string (Gemini may return dicts/lists)
    for key in narrative:
        val = narrative[key]
        if isinstance(val, list):
            narrative[key] = "\n".join(
                item if isinstance(item, str) else json.dumps(item, ensure_ascii=False)
                for item in val
            )
        elif not isinstance(val, str):
            narrative[key] = json.dumps(val, ensure_ascii=False, indent=2)

    # 2. Build Word document ---------------------------------------------------
    doc = Document()

    # Cover header
    t = doc.add_heading("CREDIT APPRAISAL MEMO (CAM)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Intelli-Credit AI  —  CONFIDENTIAL")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    # Summary info table
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = "Table Grid"
    info_table.rows[0].cells[0].text = "Company"
    info_table.rows[0].cells[1].text = request.company_name
    _table_row(info_table, "Sector", request.sector)
    _table_row(info_table, "Loan Requested", f"Rs {request.loan_requested:,.0f}")
    _table_row(
        info_table,
        "Collateral",
        f"{request.collateral_description or 'Unsecured'} "
        f"(Rs {request.collateral_value or 0:,.0f})",
    )
    _table_row(info_table, "Date of Appraisal", str(date.today()))
    _table_row(info_table, "Credit Score", f"{score.total} / 100")
    _table_row(info_table, "Decision", decision.status.upper())
    if decision.loan_amount_recommended:
        _table_row(info_table, "Recommended Loan", f"Rs {decision.loan_amount_recommended:,.0f}")
    if decision.interest_rate:
        _table_row(info_table, "Interest Rate", f"{decision.interest_rate}% p.a.")
    if decision.tenor_months:
        _table_row(info_table, "Tenor", f"{decision.tenor_months} months")
    doc.add_paragraph()

    # ── Section 1: Executive Summary
    _heading(doc, "1. Executive Summary")
    doc.add_paragraph(narrative.get("executive_summary", "N/A"))

    # ── Section 2: Applicant Profile
    _heading(doc, "2. Applicant Profile")
    doc.add_paragraph(narrative.get("applicant_profile", "N/A"))

    # ── Section 3: Five Cs
    _heading(doc, "3. Five Cs of Credit Analysis")
    five_cs_sections = [
        ("3.1  Character", score.character, 25, score_explanation.character_reasons, narrative.get("character_analysis", "")),
        ("3.2  Capacity",  score.capacity,  25, score_explanation.capacity_reasons,  narrative.get("capacity_analysis",  "")),
        ("3.3  Capital",   score.capital,   20, score_explanation.capital_reasons,   narrative.get("capital_analysis",   "")),
        ("3.4  Collateral",score.collateral,15, score_explanation.collateral_reasons,narrative.get("collateral_analysis","")),
        ("3.5  Conditions",score.conditions,15, score_explanation.conditions_reasons,narrative.get("conditions_analysis","")),
    ]
    for title, c_score, max_score, reasons, narrative_text in five_cs_sections:
        _heading(doc, f"{title}  [{c_score}/{max_score}]", level=2)
        if narrative_text:
            doc.add_paragraph(narrative_text)
        if reasons:
            doc.add_paragraph("Scoring breakdown:").runs[0].bold = True
            for reason in reasons:
                doc.add_paragraph(f"  \u2022  {reason}")
        doc.add_paragraph()

    # ── Section 4: Financial Analysis table
    _heading(doc, "4. Financial Analysis")
    if request.financial_data:
        fs = request.financial_data.financial_summary
        fr = request.financial_data.financial_ratios
        ca = request.financial_data.cashflow_analysis

        fin_table = doc.add_table(rows=1, cols=2)
        fin_table.style = "Table Grid"
        fin_table.rows[0].cells[0].text = "Metric"
        fin_table.rows[0].cells[1].text = "Value"

        fin_rows = [
            ("Revenue",             f"Rs {fs.revenue:,.0f}"             if fs.revenue is not None             else "N/A"),
            ("Net Profit",          f"Rs {fs.profit:,.0f}"              if fs.profit is not None              else "N/A"),
            ("Total Assets",        f"Rs {fs.total_assets:,.0f}"        if fs.total_assets is not None        else "N/A"),
            ("Total Liabilities",   f"Rs {fs.total_liabilities:,.0f}"   if fs.total_liabilities is not None   else "N/A"),
            ("Net Worth",           f"Rs {fs.net_worth:,.0f}"           if fs.net_worth is not None           else "N/A"),
            ("Debt-to-Equity",      f"{fr.debt_to_equity:.2f}"          if fr.debt_to_equity is not None      else "N/A"),
            ("Current Ratio",       f"{fr.current_ratio:.2f}"           if fr.current_ratio is not None       else "N/A"),
            ("Profit Margin",       f"{fr.profit_margin*100:.1f}%"      if fr.profit_margin is not None       else "N/A"),
            ("DSCR",                f"{fr.dscr:.2f}"                    if fr.dscr is not None               else "N/A"),
            ("Avg Monthly Inflow",  f"Rs {ca.monthly_avg_inflow:,.0f}"  if ca.monthly_avg_inflow is not None  else "N/A"),
            ("Avg Monthly Outflow", f"Rs {ca.monthly_avg_outflow:,.0f}" if ca.monthly_avg_outflow is not None else "N/A"),
            ("Cashflow Volatility", ca.cashflow_volatility              if ca.cashflow_volatility is not None else "N/A"),
        ]
        for label, val in fin_rows:
            _table_row(fin_table, label, val)
        doc.add_paragraph()
    else:
        doc.add_paragraph("Financial data not available.")

    # ── Section 5: Risk Matrix
    _heading(doc, "5. Risk Matrix")
    doc.add_paragraph(narrative.get("risk_matrix", "N/A"))

    if request.risk_signals:
        risk_table = doc.add_table(rows=1, cols=4)
        risk_table.style = "Table Grid"
        hdr = risk_table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "Risk Type", "Entity", "Severity", "Evidence"
        )
        for sig in request.risk_signals:
            row = risk_table.add_row()
            row.cells[0].text = sig.risk_type
            row.cells[1].text = sig.entity
            row.cells[2].text = str(sig.severity)
            row.cells[3].text = sig.evidence[:120]
        doc.add_paragraph()

    # ── Section 6: Score Summary
    _heading(doc, "6. Credit Score Summary")
    score_table = doc.add_table(rows=1, cols=3)
    score_table.style = "Table Grid"
    sh = score_table.rows[0].cells
    sh[0].text, sh[1].text, sh[2].text = "Parameter", "Score", "Max"
    for label, val, max_val in [
        ("Character",  score.character,  25),
        ("Capacity",   score.capacity,   25),
        ("Capital",    score.capital,    20),
        ("Collateral", score.collateral, 15),
        ("Conditions", score.conditions, 15),
        ("TOTAL",      score.total,      100),
    ]:
        row = score_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(val)
        row.cells[2].text = str(max_val)
    doc.add_paragraph()

    # ── Section 7: Recommendation
    _heading(doc, "7. Recommendation & Decision")
    doc.add_paragraph(narrative.get("recommendation", decision.detailed_reasoning))
    doc.add_paragraph()

    decision_para = doc.add_paragraph(f"FINAL DECISION: {decision.status.upper()}")
    decision_para.runs[0].bold = True
    decision_para.runs[0].font.size = Pt(13)
    doc.add_paragraph(decision.primary_reason)

    # 3. Save ------------------------------------------------------------------
    os.makedirs(CAM_STORAGE_DIR, exist_ok=True)
    safe_name = _safe_company_name(request.company_name)
    uid = uuid.uuid4().hex[:6]
    docx_filename = f"{safe_name}_{uid}_CAM.docx"
    docx_path = os.path.join(CAM_STORAGE_DIR, docx_filename)
    doc.save(docx_path)

    # 4. Convert to PDF --------------------------------------------------------
    pdf_filename: str | None = None
    if _DOCX2PDF_AVAILABLE:
        try:
            pdf_filename = f"{safe_name}_{uid}_CAM.pdf"
            pdf_path = os.path.join(CAM_STORAGE_DIR, pdf_filename)
            docx2pdf_convert(docx_path, pdf_path)
        except Exception as e:
            print(f"[cam_generator] PDF conversion failed: {e}")
            pdf_filename = None

    return docx_filename, pdf_filename
